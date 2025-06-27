"""DOCX to Markdown converter."""

import shutil
import subprocess
import tempfile
from typing import Any, Optional, Tuple, Union

from airweave.core.logging import logger
from airweave.platform.file_handling.conversion._base import (
    DocumentConverter,
    DocumentConverterResult,
)
from airweave.platform.sync.async_helpers import run_in_thread_pool


class DocxConverter(DocumentConverter):
    """Converts DOCX files to Markdown using python-docx, mammoth, or pandoc as fallbacks."""

    async def convert(self, local_path: str, **kwargs: Any) -> Union[None, DocumentConverterResult]:
        """Convert a DOCX file to markdown.

        Args:
            local_path: Path to the DOCX file
            **kwargs: Additional arguments passed to converters

        Returns:
            DocumentConverterResult containing the markdown text
        """
        extension = kwargs.get("file_extension", "")
        if extension.lower() != ".docx":
            return None

        md_content = ""
        title = None

        try:
            # Try using python-docx first
            try:
                md_content, title = await self._convert_with_python_docx(local_path)
            except ImportError:
                # Fall back to mammoth if python-docx is not available
                try:
                    md_content = await self._convert_with_mammoth(local_path)
                except ImportError:
                    # If neither library is available, try using pandoc
                    try:
                        md_content = await self._convert_with_pandoc(local_path)
                    except Exception as e:
                        logger.error(f"Error converting DOCX with external tools: {str(e)}")
                        return DocumentConverterResult(
                            title=None,
                            text_content=(
                                "DOCX conversion requires python-docx, mammoth, or pandoc."
                            ),
                        )
        except Exception as e:
            logger.error(f"Error converting DOCX: {str(e)}")
            return None

        return DocumentConverterResult(title=title, text_content=md_content.strip())

    async def _convert_with_python_docx(self, local_path: str) -> Tuple[str, Optional[str]]:
        """Convert DOCX using python-docx library.

        Args:
            local_path: Path to the DOCX file

        Returns:
            Tuple of (markdown_content, title)
        """
        import docx

        md_content = ""
        title = None

        doc = docx.Document(local_path)

        # Try to extract title from core properties
        if doc.core_properties.title:
            title = doc.core_properties.title

        # Process paragraphs
        for para in doc.paragraphs:
            if para.text:
                # Check if it's a heading
                if para.style.name.startswith("Heading"):
                    level = (
                        int(para.style.name.replace("Heading", ""))
                        if para.style.name != "Heading"
                        else 1
                    )
                    md_content += f"\n{'#' * level} {para.text}\n"
                else:
                    md_content += f"{para.text}\n\n"

        # Process tables
        for table in doc.tables:
            md_table = []
            # Add header row
            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            md_table.append("| " + " | ".join(header_row) + " |")
            md_table.append("|" + "|".join(["---" for _ in header_row]) + "|")

            # Add data rows
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                md_table.append("| " + " | ".join(cells) + " |")

            md_content += "\n" + "\n".join(md_table) + "\n\n"

        return md_content, title

    async def _convert_with_mammoth(self, local_path: str) -> str:
        """Convert DOCX using mammoth library.

        Args:
            local_path: Path to the DOCX file

        Returns:
            Markdown content
        """
        import mammoth

        with open(local_path, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            return result.value

    async def _convert_with_pandoc(self, local_path: str) -> str:
        """Convert DOCX file to markdown using pandoc.

        Args:
            local_path: Path to the DOCX file

        Returns:
            Markdown content
        """
        if not shutil.which("pandoc"):
            raise RuntimeError("Pandoc is not installed")

        with tempfile.NamedTemporaryFile(suffix=".md") as temp_file:
            # Run pandoc in thread pool to avoid blocking
            def _run_pandoc():
                return subprocess.run(["pandoc", local_path, "-o", temp_file.name], check=True)

            await run_in_thread_pool(_run_pandoc)

            with open(temp_file.name, "r") as f:
                return f.read()
